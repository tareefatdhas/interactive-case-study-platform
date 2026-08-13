from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
KIT = ROOT / "media-kit"
OUT = KIT / "Classfully-Brand-Guidelines.docx"

INK = "101A38"
VIOLET = "5146E5"
PAPER = "FFFEFA"
CORAL = "DF664E"
MUTED = "697087"
LINE = "E3E5ED"
SOFT = "F8F7FB"
LILAC = "D9D4FF"
SUCCESS = "3AA45A"


def rgb(value):
    return RGBColor.from_string(value)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=110, start=150, bottom=110, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size=8, edges=("top", "left", "bottom", "right")):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in edges:
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        el = borders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(total * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_font(run, name="Arial", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def add_text(doc, text, size=10.5, color=INK, bold=False, italic=False, before=0, after=7, align=None, font="Arial", keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), font, size, color, bold, italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_kicker(doc, text):
    p = add_text(doc, text.upper(), size=8.5, color=VIOLET, bold=True, after=6, keep=True)
    p.runs[0].font.letter_spacing = Pt(1.1)
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.16
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), size=10.25, bold=True)
        set_font(p.add_run(text[len(bold_lead):]), size=10.25)
    else:
        set_font(p.add_run(text), size=10.25)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.16
    set_font(p.add_run(text), size=10.25)
    return p


def add_callout(doc, label, body, fill="F0EFFF", accent=VIOLET):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_border(cell, color=accent, size=12, edges=("left",))
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(label.upper()), size=8, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.16
    set_font(p2.add_run(body), size=10.25, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_color_row(table, row_index, name, hex_code, role, light=False):
    row = table.rows[row_index]
    set_cell_fill(row.cells[0], hex_code)
    for cell in row.cells:
        set_cell_border(cell)
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(hex_code), size=8.5, color=INK if light else PAPER, bold=True)
    for idx, value in enumerate((name, role), start=1):
        p = row.cells[idx].paragraphs[0]
        set_font(p.add_run(value), size=9.5, color=INK, bold=idx == 1)


def add_page_break(doc):
    doc.add_page_break()


def add_picture(doc, source, alt_text, **kwargs):
    shape = doc.add_picture(str(source), **kwargs)
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    return shape


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.68)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.78)
section.right_margin = Inches(0.78)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = rgb(INK)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.18

for name, size, color, before, after in [
    ("Heading 1", 20, INK, 17, 8),
    ("Heading 2", 14, VIOLET, 13, 6),
    ("Heading 3", 11.5, INK, 9, 4),
]:
    style = styles[name]
    style.font.name = "Georgia" if name != "Heading 3" else "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
    style.font.size = Pt(size)
    style.font.bold = name == "Heading 3"
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Number"):
    style = styles[style_name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(10.25)
    style.paragraph_format.left_indent = Inches(0.35)
    style.paragraph_format.first_line_indent = Inches(-0.18)

header = section.header
header_p = header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header_p.add_run("CLASSFULLY  /  BRAND GUIDELINES"), size=7.5, color=MUTED, bold=True)
footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer_p.add_run("Classfully brand reference  •  August 2026"), size=7.5, color=MUTED)

# Cover
add_picture(doc, KIT / "banners/classfully-master-16x9-3840x2160.png", "Classfully response points gathering into a shared ripple beside the Classfully positioning line", width=Inches(6.94), height=Inches(3.9))
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
add_kicker(doc, "Brand reference · Version 1.0")
p = add_text(doc, "Classfully brand guidelines", size=31, color=INK, after=5, font="Georgia", keep=True)
p.paragraph_format.line_spacing = 1.0
add_text(doc, "A practical system for product, classroom, media, and partner communications.", size=13, color=MUTED, after=19)
add_callout(doc, "The idea", "Classfully makes individual student signals visible as a shared classroom response. The brand should feel intelligent, alive, calm, and unmistakably built for university teaching.")
add_text(doc, "Use this document with the files in the Classfully media kit. When a placement has its own size or crop rules, begin with the closest supplied master and preserve the safe area.", size=9.5, color=MUTED, before=12, after=0)

# Foundation
add_page_break(doc)
add_kicker(doc, "01 · Brand foundation")
add_heading(doc, "The interactive layer for university courses", 1)
add_text(doc, "Classfully works alongside an instructor's existing slides and lesson materials. It helps instructors invite participation at the right moment, understand the room, and build a useful record of student progress over time.")
add_heading(doc, "Positioning", 2)
add_callout(doc, "Primary line", "Make the classroom interactive. Build every student's journey.", fill="F8F7FB")
add_heading(doc, "Brand promise", 2)
add_text(doc, "Every student gets a clear way to participate. Every instructor gets a clearer view of the room. Every class session becomes part of a coherent course record.")
add_heading(doc, "Personality", 2)
personality = [
    ("Intelligent, not academic-sounding", "Use clear ideas and evidence without institutional jargon."),
    ("Playful, not childish", "Create anticipation and tactile delight that still feels right for adults."),
    ("Alive, not busy", "Movement should reveal participation, progress, or change."),
    ("Warm, not casual", "Sound human and encouraging without becoming cute or overfamiliar."),
    ("Confident, not grandiose", "Explain what the product helps people do. Avoid inflated claims."),
]
for lead, detail in personality:
    add_bullet(doc, f"{lead}. {detail}", f"{lead}.")
add_heading(doc, "Audience", 2)
add_text(doc, "The primary buyer and operator is a university instructor. The most frequent participant is a student using a phone. Department leaders and teaching teams are secondary audiences. Every brand decision should respect both the teaching environment and the student's limited attention.")

# Logo
add_page_break(doc)
add_kicker(doc, "02 · Identity")
add_heading(doc, "Logo and mark", 1)
add_picture(doc, KIT / "logos/classfully-lockup-color-2400.png", "Classfully logo with overlapping conversation shapes, a check form, and a coral arrival point", width=Inches(6.2))
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
add_text(doc, "The overlapping conversation shapes and check form represent contribution becoming visible. The coral point is a live arrival signal. The period in the wordmark gives the name a confident, editorial finish.", size=9.5, color=MUTED)
add_heading(doc, "Use the right version", 2)
table = doc.add_table(rows=4, cols=2)
set_table_geometry(table, [1.7, 4.8])
rows = [
    ("Color lockup", "Default for warm paper and white backgrounds."),
    ("Reversed lockup", "Use only on deep navy or an equally quiet dark surface."),
    ("Standalone mark", "Use for avatars, app icons, small navigation surfaces, and constrained squares."),
    ("Wordmark", "Keep the mark with the name in first-use and external brand contexts."),
]
for idx, (label, detail) in enumerate(rows):
    for cell in table.rows[idx].cells:
        set_cell_border(cell)
        if idx % 2:
            set_cell_fill(cell, SOFT)
    set_font(table.cell(idx, 0).paragraphs[0].add_run(label), size=9.5, bold=True)
    set_font(table.cell(idx, 1).paragraphs[0].add_run(detail), size=9.5)
add_heading(doc, "Clear space and minimum size", 2)
add_bullet(doc, "Leave at least one mark-height of clear space around the complete lockup.")
add_bullet(doc, "Do not render the complete lockup below 120 px wide on screen. Use the standalone mark when space is tighter.")
add_bullet(doc, "For print, keep the complete lockup at least 32 mm wide and the standalone mark at least 8 mm wide.")
add_heading(doc, "Never", 2)
for item in ["Stretch, skew, rotate, or redraw the mark.", "Change individual logo colors or remove the coral signal point.", "Add bevels, drop shadows, outlines, or animation to the external logo.", "Place the logo over a busy part of the ripple field.", "Use an AI-generated approximation of the logo or wordmark."]:
    add_bullet(doc, item)

# Color
add_page_break(doc)
add_kicker(doc, "03 · Visual system")
add_heading(doc, "Color", 1)
add_text(doc, "Classfully is primarily warm paper, deep navy, and violet. Coral is a signal, not a second primary color. Supporting colors should identify meaning in a data display rather than decorate a screen.")
table = doc.add_table(rows=7, cols=3)
set_table_geometry(table, [1.25, 1.55, 3.7])
colors = [
    ("Warm paper", PAPER, "Primary canvas and calm negative space", True),
    ("Deep navy", INK, "Primary text, dark surfaces, and strong actions", False),
    ("Classfully violet", VIOLET, "Brand action, live interaction, and participation", False),
    ("Arrival coral", CORAL, "One meaningful arrival, emphasis, or alert", False),
    ("Muted slate", MUTED, "Secondary text and explanatory labels", False),
    ("Soft lilac", LILAC, "Quiet selection or participation background", True),
    ("Success green", SUCCESS, "Connected, complete, or confirmed states", False),
]
for i, (name, code, role, light) in enumerate(colors):
    add_color_row(table, i, name, code, role, light)
add_heading(doc, "Color hierarchy", 2)
add_number(doc, "Begin with warm paper and deep navy so the classroom surface stays readable.")
add_number(doc, "Use violet for the current activity, selected state, or response field.")
add_number(doc, "Introduce coral only when it points to a meaningful moment.")
add_number(doc, "Use semantic colors consistently. Do not use color as the only carrier of meaning.")
add_callout(doc, "Avoid gradient slop", "Do not use broad decorative gradients, glowing blobs, or arbitrary color washes. Depth should come from spacing, restrained shadows, tangible button layers, and the response field itself.", fill="FFF4F1", accent=CORAL)

# Typography
add_page_break(doc)
add_kicker(doc, "04 · Typography")
add_heading(doc, "Editorial clarity meets usable interface type", 1)
add_text(doc, "Classfully pairs Fraunces with Inter. The display face gives important classroom moments a human, editorial voice. The interface face keeps controls, data, and instructions fast to read.")
table = doc.add_table(rows=3, cols=3)
set_table_geometry(table, [1.45, 1.7, 3.35])
font_rows = [
    ("Display", "Fraunces", "Large questions, key insights, section titles, and major numbers"),
    ("Interface", "Inter", "Navigation, buttons, labels, instructional copy, and data"),
    ("Portable fallback", "Georgia + Arial", "Media exports and documents where product fonts cannot travel"),
]
for i, row in enumerate(font_rows):
    for j, value in enumerate(row):
        cell = table.cell(i, j)
        set_cell_border(cell)
        if i % 2:
            set_cell_fill(cell, SOFT)
        set_font(cell.paragraphs[0].add_run(value), name="Georgia" if j == 1 and i != 1 else "Arial", size=9.5, bold=j < 2)
add_heading(doc, "Type roles", 2)
add_bullet(doc, "Display headings: sentence case, generous line height, and a maximum of roughly 10 to 12 words when projected.", "Display headings:")
add_bullet(doc, "UI labels: concise and concrete. Prefer “Start class” over “Initialize session.”", "UI labels:")
add_bullet(doc, "Eyebrows: short uppercase labels with restrained tracking, used for category and status.", "Eyebrows:")
add_bullet(doc, "Numbers: use tabular figures for live counts, scores, timers, and join codes.", "Numbers:")
add_bullet(doc, "Emphasis: prefer hierarchy, weight, and spacing. Do not rely on excessive bold, all caps, or colored text.", "Emphasis:")
add_callout(doc, "Projection check", "If an instructor or student has to reread a projected question, the type is too small, too dense, or too decorative.")

# Imagery + ripple
add_page_break(doc)
add_kicker(doc, "05 · Signature visual language")
add_heading(doc, "The response ripple", 1)
add_picture(doc, KIT / "source/classfully-response-field-master.png", "Indigo student response points arriving through curved paths and settling into an organic shared ripple", width=Inches(5.7))
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
add_text(doc, "Individual points arrive, gather, and shape a shared surface. This is the visual expression of Classfully: a room becoming legible without reducing students to a spreadsheet.", size=9.5, color=MUTED)
add_heading(doc, "Meaning before decoration", 2)
add_bullet(doc, "A point represents a response, contribution, learning moment, or arrival.")
add_bullet(doc, "A path represents movement from an individual device into the classroom signal.")
add_bullet(doc, "A ripple represents collective effect, accumulation, or progress over time.")
add_bullet(doc, "A coral point marks one meaningful arrival or focal moment. It should remain rare.")
add_heading(doc, "When to use it", 2)
add_text(doc, "Use the response field in presentation transitions, post-response confirmation, student progress, launch artwork, product listings, and calm waiting states. Do not use it behind dense forms, detailed tables, long policy copy, or every card in the interface.")
add_heading(doc, "Imagery principles", 2)
add_bullet(doc, "Prefer real product views, real classroom context, and abstract response imagery over generic stock photography.")
add_bullet(doc, "Keep people and devices honest. Do not fabricate product capabilities in promotional mockups.")
add_bullet(doc, "Use warm, natural light and real texture. Avoid glossy 3D blobs, neon effects, confetti, and cartoon classroom imagery.")

# Motion
add_page_break(doc)
add_kicker(doc, "06 · Motion and interaction")
add_heading(doc, "Motion should help the room feel something happen", 1)
add_text(doc, "Classfully motion connects an action on a student's phone with a visible change in the classroom. It should acknowledge, transport, settle, and reveal. It should never delay the instructor or obscure the data.")
add_heading(doc, "Core motion sequence", 2)
motion = [
    ("Acknowledge", "The selected control compresses or responds immediately."),
    ("Transport", "A signal appears to leave the student's action and move into the shared system."),
    ("Settle", "The incoming point lands into the relevant category, cluster, or trend."),
    ("Reveal", "Counts, labels, or aggregate understanding update after the motion makes the change understandable."),
]
for lead, detail in motion:
    add_number(doc, f"{lead}. {detail}")
add_heading(doc, "Timing", 2)
add_bullet(doc, "Micro feedback: about 120 to 220 ms.")
add_bullet(doc, "Response transport: about 650 to 1,200 ms, paced by actual incoming responses.")
add_bullet(doc, "View transition: about 280 to 450 ms.")
add_bullet(doc, "Ambient waiting motion: slow, low-amplitude, and never the only sign of state.")
add_heading(doc, "Motion rules", 2)
add_bullet(doc, "Responses should arrive at the pace students submit them. Do not create a theatrical queue that disconnects the visualization from the room.")
add_bullet(doc, "Use spring and settle behavior for tactile controls, but keep labels stable and readable.")
add_bullet(doc, "Respect reduced-motion preferences with an immediate state change and restrained fade.")
add_bullet(doc, "Do not promise phone haptics where the browser cannot reliably provide them. Visual feedback must stand on its own.")
add_callout(doc, "Purpose test", "If removing an animation does not reduce understanding, reassurance, anticipation, or delight at a meaningful moment, the animation probably does not belong.")

# Product UI
add_page_break(doc)
add_kicker(doc, "07 · Product interface")
add_heading(doc, "Different surfaces, one system", 1)
add_text(doc, "Classfully has three live surfaces with different jobs. Consistency comes from shared type, color, controls, and motion. Density should change according to the audience and viewing distance.")
table = doc.add_table(rows=4, cols=3)
set_table_geometry(table, [1.45, 2.35, 2.7])
headers = ("Surface", "Primary job", "Design priority")
for j, value in enumerate(headers):
    cell = table.cell(0, j)
    set_cell_fill(cell, INK)
    set_cell_border(cell, color=INK)
    set_font(cell.paragraphs[0].add_run(value), size=9, color=PAPER, bold=True)
surface_rows = [
    ("Instructor console", "Control the current moment and see what needs attention", "Now, next, questions, and safe recovery actions"),
    ("Classroom display", "Make the prompt and collective response legible across the room", "Large type, low clutter, stable data, and visible join access"),
    ("Student phone", "Respond quickly and understand what happens next", "Thumb reach, tactile controls, privacy clarity, and resilient connection"),
]
for i, row in enumerate(surface_rows, start=1):
    for j, value in enumerate(row):
        cell = table.cell(i, j)
        set_cell_border(cell)
        if i % 2 == 0:
            set_cell_fill(cell, SOFT)
        set_font(cell.paragraphs[0].add_run(value), size=9.25, bold=j == 0)
add_heading(doc, "Components", 2)
add_bullet(doc, "Buttons: keep the tactile layered effect for meaningful student actions. Do not add a competing pointer-focus ring. Preserve a clear keyboard focus state.", "Buttons:")
add_bullet(doc, "Cards: use borders and spacing before shadow, tint, or gradient. A card should group related information, not decorate empty space.", "Cards:")
add_bullet(doc, "Status: say what is true and useful. “Responses are open” is better than a vague “Live.”", "Status:")
add_bullet(doc, "Empty states: explain what will appear and what action starts it. Never flash fictional classroom data while real data loads.", "Empty states:")
add_bullet(doc, "Errors: explain what happened, whether work is safe, and what the person can do next. Recover automatically where possible.", "Errors:")
add_heading(doc, "Visual restraint", 2)
add_text(doc, "Use one dominant focal action per moment. Hide or group secondary controls. Avoid stacking borders, shadows, tints, gradients, badges, and animated texture on the same component.")

# Voice
add_page_break(doc)
add_kicker(doc, "08 · Voice and copy")
add_heading(doc, "Natural language for a real classroom", 1)
add_text(doc, "Classfully copy should sound like a thoughtful instructor or teaching assistant: direct, calm, friendly, and specific. It should not sound like enterprise software or synthetic marketing copy.")
add_heading(doc, "Writing principles", 2)
principles = [
    ("Say what happens", "“Post question” instead of “Submit engagement.”"),
    ("Use classroom language", "Class, session, activity, question, response, participation, progress, and instructor."),
    ("Keep instructions short", "One action per sentence whenever possible."),
    ("Respect student agency", "Explain privacy, anonymity, scoring, and what the projector shows."),
    ("Write calm errors", "Avoid blame and give a recovery path."),
]
for lead, detail in principles:
    add_bullet(doc, f"{lead}. {detail}", f"{lead}.")
add_heading(doc, "Anti-slop rules", 2)
for item in [
    "Do not use em dashes. Use a period, comma, colon, or parentheses when needed.",
    "Avoid “seamless,” “unlock,” “empower,” “revolutionize,” “ecosystem,” “journey” as filler, and other generic product language.",
    "Do not use rhetorical fragments simply to sound dramatic.",
    "Do not overuse paired constructions such as “not just X, but Y.”",
    "Do not repeat the same benefit in a headline, subhead, and button.",
    "Avoid emojis in core product copy. Icons already carry visual tone.",
]:
    add_bullet(doc, item)
add_heading(doc, "Examples", 2)
table = doc.add_table(rows=5, cols=2)
set_table_geometry(table, [3.25, 3.25])
for j, value in enumerate(("Prefer", "Avoid")):
    set_cell_fill(table.cell(0, j), VIOLET if j == 0 else CORAL)
    set_cell_border(table.cell(0, j), color=VIOLET if j == 0 else CORAL)
    set_font(table.cell(0, j).paragraphs[0].add_run(value), size=9, color=PAPER, bold=True)
examples = [
    ("Start class", "Launch your live learning experience"),
    ("While the room responds", "Awaiting participant engagement"),
    ("Your response is in", "Submission successfully processed"),
    ("Ask without interrupting", "Unlock seamless classroom dialogue"),
]
for i, row in enumerate(examples, start=1):
    for j, value in enumerate(row):
        cell = table.cell(i, j)
        set_cell_border(cell)
        if i % 2 == 0:
            set_cell_fill(cell, SOFT)
        set_font(cell.paragraphs[0].add_run(value), size=9.25, color=INK, bold=j == 0)

# Accessibility + checklist
add_page_break(doc)
add_kicker(doc, "09 · Accessibility and quality")
add_heading(doc, "Designed for the room, not the ideal device", 1)
add_text(doc, "Classfully must work for a student on an older phone, an instructor under time pressure, and a projector viewed from the back of a large classroom. Accessibility is part of the brand promise because participation only works when people can actually participate.")
add_heading(doc, "Minimum expectations", 2)
for item in [
    "Meet WCAG AA contrast for interface text and essential controls.",
    "Never communicate correctness, urgency, or selection through color alone.",
    "Keep touch targets at least 44 × 44 CSS pixels and leave enough space between adjacent actions.",
    "Provide a visible keyboard focus state without showing pointer focus rings after taps.",
    "Support text enlargement and narrow mobile screens without clipping actions or footer content.",
    "Provide reduced-motion behavior that preserves clear acknowledgment and state change.",
    "Use plain language for privacy, scoring, attendance, and anonymous questions.",
    "Keep projector copy large, short, and high contrast. Test from realistic viewing distance.",
]:
    add_bullet(doc, item)
add_heading(doc, "Final review checklist", 2)
for item in [
    "Does this feel appropriate for university students and instructors?",
    "Is the most important action obvious within three seconds?",
    "Does every color, layer, shadow, and animation have a job?",
    "Are real states shown instead of placeholder classroom data?",
    "Can the interface recover without asking a room of students to refresh?",
    "Does the copy sound natural when read aloud by an instructor?",
    "Have we avoided em dashes, inflated language, decorative gradients, and generic AI styling?",
    "Does the result remain clear with motion reduced and color perception limited?",
]:
    add_bullet(doc, f"☐ {item}")
add_callout(doc, "The standard", "The product should feel special because it makes classroom participation tangible, not because every surface is embellished.", fill="F0EFFF")

# Assets
add_page_break(doc)
add_kicker(doc, "10 · Asset reference")
add_heading(doc, "Start with the media kit", 1)
add_text(doc, "Use the supplied assets before recreating a logo, banner, icon, or response field. SVG logo files are the master for print and resizing. PNG exports are prepared for common digital placements.")
asset_table = doc.add_table(rows=1, cols=3)
set_table_geometry(asset_table, [2.2, 2.8, 1.5])
for j, value in enumerate(("Use", "File", "Size")):
    cell = asset_table.cell(0, j)
    set_cell_fill(cell, INK)
    set_cell_border(cell, color=INK)
    set_font(cell.paragraphs[0].add_run(value), size=8.5, color=PAPER, bold=True)
assets = [
    ("Primary logo", "logos/classfully-lockup-color.svg", "Vector"),
    ("Profile image", "logos/classfully-profile-400.png", "400 × 400"),
    ("App icon", "logos/classfully-app-icon-1024.png", "1024 × 1024"),
    ("Universal banner", "banners/classfully-master-16x9-3840x2160.png", "3840 × 2160"),
    ("LinkedIn cover", "banners/classfully-linkedin-cover-4200x700.png", "4200 × 700"),
    ("X header", "banners/classfully-x-header-1500x500.png", "1500 × 500"),
    ("YouTube art", "banners/classfully-youtube-channel-2560x1440.png", "2560 × 1440"),
    ("Link preview", "social/classfully-open-graph-1200x630.png", "1200 × 630"),
    ("Square social", "social/classfully-square-2160x2160.png", "2160 × 2160"),
    ("Portrait social", "social/classfully-portrait-2160x2700.png", "2160 × 2700"),
    ("Vertical story", "social/classfully-story-2160x3840.png", "2160 × 3840"),
    ("Play feature", "store/classfully-google-play-feature-1024x500.png", "1024 × 500"),
]
for i, row_data in enumerate(assets, start=1):
    cells = asset_table.add_row().cells
    for j, value in enumerate(row_data):
        set_cell_border(cells[j])
        if i % 2 == 0:
            set_cell_fill(cells[j], SOFT)
        set_font(cells[j].paragraphs[0].add_run(value), size=8.25, color=INK, bold=j == 0)
add_heading(doc, "Source of truth", 2)
add_text(doc, "Media kit folder: case-study-platform/media-kit", size=10, bold=True)
add_text(doc, "Asset manifest: media-kit/asset-manifest.json", size=9.5, color=MUTED)
add_text(doc, "Regeneration script: scripts/build-media-kit.mjs", size=9.5, color=MUTED)
add_text(doc, "Before a major launch, verify the destination's current crop, file-size, and transparency requirements. Platform specifications can change.", size=9.5, color=MUTED, italic=True)

doc.core_properties.title = "Classfully Brand Guidelines"
doc.core_properties.subject = "Brand identity, voice, product interface, motion, and media guidance"
doc.core_properties.author = "Classfully"
doc.core_properties.keywords = "Classfully, brand, guidelines, university classroom, interaction"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
